from flask import Flask, make_response, render_template, request, jsonify, send_from_directory, send_file, g
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
import os, json, uuid, secrets, time, re
from datetime import datetime, timedelta
import jwt
from functools import wraps
from PIL import Image
import io
import requests as http_requests
import logging
logging.basicConfig(level=logging.WARNING)

app = Flask(__name__)
app.config['TEMPLATES_AUTO_RELOAD'] = True
JWT_EXPIRATION_HOURS = 24 * 7
UPLOAD_MAX_SIZE = 10 * 1024 * 1024   # 10MB

# ---- 请求日志：捕获所有请求 ----
@app.before_request
def log_request():
    import logging
    if request.path.startswith('/api/'):
        logging.warning(f'REQ: {request.method} {request.path} ct={request.content_type} cl={request.content_length} ua={request.headers.get("User-Agent","?")[:60]}')

# ---- 全局错误处理：所有非 JSON 错误统一返回 JSON ----
@app.errorhandler(400)
def bad_request(e):
    import logging
    logging.warning(f'400_ERROR: path={request.path} method={request.method}')
    return jsonify({"status": "error", "message": str(e), "code": 400}), 400, {"Content-Type": "application/json"}

@app.errorhandler(404)
def not_found(e):
    return jsonify({"status": "error", "message": "资源不存在", "code": 404}), 404, {"Content-Type": "application/json"}

@app.errorhandler(405)
def method_not_allowed(e):
    return jsonify({"status": "error", "message": str(e), "code": 405}), 405, {"Content-Type": "application/json"}

@app.errorhandler(500)
def server_error(e):
    return jsonify({"status": "error", "message": "服务器内部错误", "code": 500}), 500, {"Content-Type": "application/json"}

# 统一 JSON 响应编码，确保所有 API 响应使用 UTF-8
@app.after_request
def ensure_utf8(response):
    if response.content_type and 'application/json' in response.content_type:
        response.headers['Content-Type'] = 'application/json; charset=utf-8'
    return response

# ---- 持久化 SECRET_KEY ----
SECRET_FILE = ".secret_key"
if os.path.exists(SECRET_FILE):
    with open(SECRET_FILE, "r") as f:
        app.config['SECRET_KEY'] = f.read().strip()
else:
    app.config['SECRET_KEY'] = secrets.token_hex(32)
    with open(SECRET_FILE, "w") as f:
        f.write(app.config['SECRET_KEY'])

# ---- 内存缓存 ----
_cache = {}
def cache_get(key, max_age=5):
    """取缓存，过期返回None"""
    entry = _cache.get(key)
    if entry and time.time() - entry['ts'] < max_age:
        return entry['val']
    return None

def cache_set(key, val):
    _cache[key] = {'ts': time.time(), 'val': val}

def cache_del(key):
    _cache.pop(key, None)

# ---- 登录速率限制 ----
_login_attempts = {}
LOGIN_MAX_ATTEMPTS = 5
LOGIN_WINDOW = 300  # 5 分钟内

def check_login_rate(username):
    now = time.time()
    attempts = _login_attempts.get(username, [])
    attempts = [t for t in attempts if now - t < LOGIN_WINDOW]
    if len(attempts) >= LOGIN_MAX_ATTEMPTS:
        return False, int(LOGIN_WINDOW - (now - attempts[0]))
    attempts.append(now)
    _login_attempts[username] = attempts
    # 定期清理，防止内存泄漏
    if len(_login_attempts) > 200:
        _login_attempts.clear()
    return True, 0

def is_valid_username(username):
    """用户名：2-20位，允许字母、数字、下划线、中文"""
    return bool(re.match(r'^[a-zA-Z0-9_\u4e00-\u9fff]{2,20}$', str(username)))

DOCS_DIR = "docs"
USERS_DIR = "users"
COMMENTS_DIR = "comments"
AVATARS_DIR = os.path.join(USERS_DIR, "avatars")
COVERS_DIR = "static/covers"
UPLOADS_DIR = "static/uploads"
USERS_FILE = os.path.join(USERS_DIR, "users.json")
COUNTER_FILE = os.path.join(DOCS_DIR, ".counter.txt")
MESSAGES_FILE = os.path.join(USERS_DIR, "messages.json")
COLLECTIONS_DIR = "collections"

os.makedirs(DOCS_DIR, exist_ok=True)
os.makedirs(AVATARS_DIR, exist_ok=True)
os.makedirs(COMMENTS_DIR, exist_ok=True)
os.makedirs(UPLOADS_DIR, exist_ok=True)
os.makedirs(COVERS_DIR, exist_ok=True)
os.makedirs(COLLECTIONS_DIR, exist_ok=True)

STICKERS_DIR = os.path.join(UPLOADS_DIR, "stickers")
os.makedirs(STICKERS_DIR, exist_ok=True)

VIDEOS_DIR = os.path.join(UPLOADS_DIR, "videos")
os.makedirs(VIDEOS_DIR, exist_ok=True)

# ---- 工具函数 ----
def allowed_ext():
    return {'png', 'jpg', 'jpeg', 'gif', 'webp'}

def allowed_video_ext():
    return {'mp4', 'webm', 'mov', 'avi'}

def is_allowed_ext(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_ext()

def make_thumb(image_path, thumb_path, max_size=300):
    """生成缩略图（WebP）"""
    try:
        img = Image.open(image_path)
        img.thumbnail((max_size, max_size))
        img.save(thumb_path, 'WEBP', quality=80)
    except Exception:
        pass  # 缩略图失败不阻塞主流程

def optimize_image(image_path):
    """无损压缩图片（重写覆盖原图）"""
    try:
        img = Image.open(image_path)
        fmt = img.format or 'JPEG'
        if fmt in ('PNG',):
            img.save(image_path, 'PNG', optimize=True)
        elif fmt in ('JPEG', 'JPG'):
            img.save(image_path, 'JPEG', quality=85, optimize=True)
    except Exception:
        pass

# ---------- JWT 工具 ----------
def create_token(username):
    users = load_users()
    role = users.get(username, {}).get('role', 'user')
    payload = {
        'username': username,
        'role': role,
        'exp': datetime.utcnow() + timedelta(hours=JWT_EXPIRATION_HOURS)
    }
    return jwt.encode(payload, app.config['SECRET_KEY'], algorithm='HS256')

def verify_token(token):
    try:
        return jwt.decode(token, app.config['SECRET_KEY'], algorithms=['HS256'])
    except Exception:
        return None

def login_required(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        token = request.headers.get('Authorization', '').replace('Bearer ', '')
        if not token:
            return jsonify({'status': 'error', 'message': '未登录'}), 401
        payload = verify_token(token)
        if not payload:
            return jsonify({'status': 'error', 'message': '登录已过期'}), 401
        g.username = payload['username']
        g.role = payload.get('role', 'user')
        return f(*args, **kwargs)
    return decorated

def get_optional_user():
    """尝试从 token 获取用户，但不强制登录"""
    token = request.headers.get('Authorization', '').replace('Bearer ', '')
    if token:
        payload = verify_token(token)
        if payload:
            g.username = payload['username']
            g.role = payload.get('role', 'user')
            return
    g.username = None
    g.role = 'user'

# ---------- 文件读写 ----------
def load_users():
    cached = cache_get('users', max_age=5)
    if cached is not None:
        return cached
    if not os.path.exists(USERS_FILE):
        return {}
    with open(USERS_FILE, "r", encoding="utf-8") as f:
        users = json.load(f)
    for u in users.values():
        if "role" not in u:
            u["role"] = "user"
    if "114514" in users:
        users["114514"]["role"] = "owner"
    cache_set('users', users)
    return users

def save_users(users):
    try:
        with open(USERS_FILE, "w", encoding="utf-8") as f:
            json.dump(users, f, ensure_ascii=False, indent=2)
    except Exception:
        pass
    cache_del('users')  # 让缓存失效

def add_notification(username, noti):
    users = load_users()
    if username not in users:
        return
    if 'notifications' not in users[username]:
        users[username]['notifications'] = []
    new_noti = {
        **noti,
        'id': str(uuid.uuid4()),
        'read': False,
        'timestamp': datetime.now().isoformat()
    }
    users[username]['notifications'].insert(0, new_noti)
    if len(users[username]['notifications']) > 50:
        users[username]['notifications'] = users[username]['notifications'][:50]
    save_users(users)

def get_meta_path(doc_id):
    return os.path.join(DOCS_DIR, f"{doc_id}.meta.json")

def load_meta(doc_id):
    path = get_meta_path(doc_id)
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    return {"authors": [], "views": 0, "likes": 0, "liked_by": [], "edit_mode": "public", "allowed_users": [], "title": doc_id, "status": "published", "text_stroke": False, "cover_thumb": None}

def save_meta(doc_id, meta):
    path = get_meta_path(doc_id)
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(meta, f, ensure_ascii=False, indent=2)
    except Exception:
        pass
    cache_del('docs_list_updated')
    cache_del('docs_list_number')
    cache_del('docs_list_views')

def check_edit_permission(doc_id, username):
    if not username:
        return False
    meta = load_meta(doc_id)
    mode = meta.get('edit_mode', 'public')
    if mode == 'public':
        return True
    users = load_users()
    nickname = users.get(username, {}).get('nickname', username)
    if mode == 'authors':
        return nickname in meta.get('authors', [])
    if mode == 'restricted':
        return username in meta.get('allowed_users', [])
    return False

def get_next_number():
    import fcntl
    if not os.path.exists(COUNTER_FILE):
        with open(COUNTER_FILE, "w") as f:
            fcntl.flock(f, fcntl.LOCK_EX)
            f.write("1")
            fcntl.flock(f, fcntl.LOCK_UN)
        return 1
    with open(COUNTER_FILE, "r+") as f:
        fcntl.flock(f, fcntl.LOCK_EX)
        num = int(f.read().strip())
        f.seek(0)
        f.write(str(num + 1))
        f.truncate()
        fcntl.flock(f, fcntl.LOCK_UN)
        return num

def get_comments_path(doc_id):
    return os.path.join(COMMENTS_DIR, f"{doc_id}.comments.json")

def load_comments(doc_id):
    cached = cache_get(f'comments_{doc_id}', max_age=10)
    if cached is not None:
        return cached
    path = get_comments_path(doc_id)
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
            cache_set(f'comments_{doc_id}', data)
            return data
    return []

def save_comments(doc_id, comments):
    path = get_comments_path(doc_id)
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(comments, f, ensure_ascii=False, indent=2)
    except Exception:
        pass
    cache_del(f'comments_{doc_id}')

# ---------- 私信存储 ----------
def load_messages():
    cached = cache_get('messages', max_age=5)
    if cached is not None:
        return cached
    if not os.path.exists(MESSAGES_FILE):
        return []
    with open(MESSAGES_FILE, "r", encoding="utf-8") as f:
        try:
            data = json.load(f)
            cache_set('messages', data)
            return data
        except Exception:
            return []

def save_messages(msgs):
    try:
        with open(MESSAGES_FILE, "w", encoding="utf-8") as f:
            json.dump(msgs, f, ensure_ascii=False, indent=2)
    except Exception:
        pass
    cache_del('messages')

# ---------- 页面路由 ----------
@app.route("/")
def index():
    resp = make_response(render_template("index.html"))
    resp.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    resp.headers["Pragma"] = "no-cache"
    resp.headers["Expires"] = "0"
    return resp

@app.route("/profile")
def profile():
    return render_template("profile.html")

# 大纲系统已删除

@app.route("/edit/<doc_id>")
def edit_page(doc_id):
    return render_template("edit.html", doc_id=doc_id)

@app.route("/read/<doc_id>")
def read_page(doc_id):
    return render_template("read.html", doc_id=doc_id)

@app.route("/user/<path:username>")
def user_page(username):
    return render_template("user.html", username=username)

# ---------- 认证 ----------
@app.route("/api/register", methods=["POST"])
def register():
    # 兼容 JSON 和表单两种 Content-Type
    if request.content_type and 'application/json' in request.content_type:
        data = request.get_json(force=True, silent=True) or {}
    else:
        data = request.form.to_dict() or {}
    username = data.get("username")
    password = data.get("password")
    nickname = data.get("nickname")
    if not username or not password or not nickname:
        return jsonify({"status": "error", "message": "缺少信息"}), 400
    if not is_valid_username(username):
        return jsonify({"status": "error", "message": "用户名须为2-20位字母、数字、下划线或中文"}), 400
    users = load_users()
    if username in users:
        return jsonify({"status": "error", "message": "用户名已存在"}), 400
    users[username] = {
        "password": generate_password_hash(password),
        "nickname": nickname,
        "avatar": None,
        "notifications": [],
        "following": [],
        "followers": []
    }
    save_users(users)
    token = create_token(username)
    return jsonify({
        "status": "ok",
        "token": token,
        "user": {
            "username": username,
            "nickname": nickname,
            "avatar": None,
            "unread": 0,
            "role": "user"
        }
    })

@app.route("/collections")
def collections_page():
    return render_template("collections.html")

@app.route("/collections/<collection_id>")
def collection_page(collection_id):
    return render_template("collection.html", collection_id=collection_id)

@app.route("/messages")
@app.route("/messages/<path:username>")
def messages_page(username=None):
    """站内信页面，可选择预设对话对象"""
    return render_template("messages.html")
@app.route("/api/login", methods=["POST"])

def login():
    # 兼容 JSON 和表单两种 Content-Type
    if request.content_type and 'application/json' in request.content_type:
        data = request.get_json(force=True, silent=True) or {}
    else:
        data = request.form.to_dict() or {}
    username = data.get("username")
    password = data.get("password")
    if not username or not password:
        return jsonify({"status": "error", "message": "用户名或密码错误"}), 401
    allowed, retry = check_login_rate(username)
    if not allowed:
        return jsonify({"status": "error", "message": f"尝试次数过多，请 {retry} 秒后再试"}), 429
    users = load_users()
    if username not in users:
        return jsonify({"status": "error", "message": "用户名或密码错误"}), 401
    if not check_password_hash(users[username]["password"], password):
        return jsonify({"status": "error", "message": "用户名或密码错误"}), 401
    token = create_token(username)
    unread = sum(1 for n in users[username].get('notifications', []) if not n.get('read', False))
    return jsonify({
        "status": "ok",
        "token": token,
        "user": {
            "username": username,
            "nickname": users[username]["nickname"],
            "avatar": users[username].get("avatar"),
            "unread": unread,
            "role": users[username].get("role", "user"),
            "is_dojin": users[username].get("dojin", False)
        }
    })

@app.route("/api/current_user", methods=["GET"])
@login_required
def current_user():
    users = load_users()
    username = g.username
    if username in users:
        return jsonify({"user": {
            "role": users[username].get("role", "user"),
            "is_dojin": users[username].get("dojin", False),
            "username": username,
            "nickname": users[username]["nickname"],
            "avatar": users[username].get("avatar")
        }})
    return jsonify({"user": None})

# ---------- 合集 ----------
def load_collections():
    """加载所有合集"""
    collections = {}
    try:
        for fname in os.listdir(COLLECTIONS_DIR):
            if fname.endswith(".json"):
                cid = fname.replace(".json", "")
                with open(os.path.join(COLLECTIONS_DIR, fname), "r", encoding="utf-8") as f:
                    collections[cid] = json.load(f)
    except OSError:
        pass
    return collections

def save_collection(collection_id, data):
    path = os.path.join(COLLECTIONS_DIR, f"{collection_id}.json")
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

def delete_collection_file(collection_id):
    path = os.path.join(COLLECTIONS_DIR, f"{collection_id}.json")
    if os.path.exists(path):
        os.remove(path)

@app.route("/api/collections", methods=["GET"])
def list_collections():
    collections = load_collections()
    # 从 header 解析当前用户名和角色
    current_username = None
    current_role = None
    auth = request.headers.get("Authorization", "")
    if auth.startswith("Bearer "):
        try:
            payload = jwt.decode(auth[7:], app.config['SECRET_KEY'], algorithms=['HS256'])
            current_username = payload.get('username')
            current_role = payload.get('role')
        except Exception:
            pass
    result = []
    for cid, cdata in collections.items():
        editors = cdata.get("editors", [])
        result.append({
            "id": cid,
            "name": cdata.get("name", "未命名合集"),
            "description": cdata.get("description", ""),
            "cover": cdata.get("cover", ""),
            "author": cdata.get("author", ""),
            "author_nickname": cdata.get("author_nickname", ""),
            "created_at": cdata.get("created_at", ""),
            "article_count": len(cdata.get("articles", [])),
            "editors": editors,
            "can_edit": can_edit_collection(cdata, current_username, current_role) if current_username else False
        })
    result.sort(key=lambda x: x["created_at"], reverse=True)
    resp = jsonify(result)
    resp.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    return resp

def can_edit_collection(cdata, username, role):
    """判断用户是否有权编辑合集"""
    editors = cdata.get("editors", [])
    return cdata.get("author") == username or username in editors or role in ("owner", "ai")


@app.route("/api/collections", methods=["POST"])
@login_required
def create_collection():
    data = request.json
    name = data.get("name", "").strip()
    description = data.get("description", "").strip()
    if not name:
        return jsonify({"status": "error", "message": "请填写合集名称"}), 400
    
    collection_id = str(uuid.uuid4())[:8]
    users = load_users()
    user_info = users.get(g.username, {})
    collection_data = {
        "name": name,
        "description": description,
        "cover": "",
        "author": g.username,
        "author_nickname": user_info.get("nickname", g.username),
        "created_at": datetime.now().isoformat(),
        "articles": [],
        "editors": []
    }
    save_collection(collection_id, collection_data)
    return jsonify({"status": "ok", "collection_id": collection_id, "collection": collection_data})

@app.route("/api/collections/<collection_id>", methods=["GET"])
def get_collection(collection_id):
    collections = load_collections()
    cdata = collections.get(collection_id)
    if not cdata:
        return jsonify({"status": "error", "message": "合集不存在"}), 404
    
    articles = []
    for doc_id in cdata.get("articles", []):
        meta = load_meta(doc_id)
        doc_path = os.path.join(DOCS_DIR, f"{doc_id}.txt")
        is_published = meta.get("status", "published") == "published"
        if os.path.exists(doc_path):
            articles.append({
                "id": doc_id,
                "title": meta.get("title", doc_id),
                "number": meta.get("number", 0),
                "authors": meta.get("authors", []),
                "cover": meta.get("cover", ""),
                "cover_thumb": meta.get("cover_thumb", ""),
                "status": meta.get("status", "published"),
                "views": meta.get("views", 0),
                "updated_at": meta.get("updated_at", "")
            })
    
    editors = cdata.get("editors", [])
    users = load_users()
    editor_nicknames = {}
    for uname in editors:
        uinfo = users.get(uname, {})
        editor_nicknames[uname] = uinfo.get("nickname", uname)

    result = {
        "id": collection_id,
        "name": cdata["name"],
        "description": cdata.get("description", ""),
        "cover": cdata.get("cover", ""),
        "author": cdata.get("author", ""),
        "author_nickname": cdata.get("author_nickname", ""),
        "created_at": cdata.get("created_at", ""),
        "articles": articles,
        "editors": editors,
        "editor_nicknames": editor_nicknames
    }
    return jsonify(result)

@app.route("/api/collections/<collection_id>", methods=["POST"])
@login_required
def update_collection(collection_id):
    collections = load_collections()
    cdata = collections.get(collection_id)
    if not cdata:
        return jsonify({"status": "error", "message": "合集不存在"}), 404
    if not can_edit_collection(cdata, g.username, g.role):
        return jsonify({"status": "error", "message": "无权修改"}), 403
    
    data = request.json
    if "name" in data:
        cdata["name"] = data["name"].strip()
    if "description" in data:
        cdata["description"] = data["description"].strip()
    if "cover" in data:
        cdata["cover"] = data["cover"]
    
    save_collection(collection_id, cdata)
    return jsonify({"status": "ok", "collection": cdata})

@app.route("/api/collections/<collection_id>", methods=["DELETE"])
@login_required
def delete_collection(collection_id):
    collections = load_collections()
    cdata = collections.get(collection_id)
    if not cdata:
        return jsonify({"status": "error", "message": "合集不存在"}), 404
    if not can_edit_collection(cdata, g.username, g.role):
        return jsonify({"status": "error", "message": "无权删除"}), 403
    
    delete_collection_file(collection_id)
    return jsonify({"status": "ok"})

@app.route("/api/collections/<collection_id>/articles", methods=["POST"])
@login_required
def add_article_to_collection(collection_id):
    collections = load_collections()
    cdata = collections.get(collection_id)
    if not cdata:
        return jsonify({"status": "error", "message": "合集不存在"}), 404
    if not can_edit_collection(cdata, g.username, g.role):
        return jsonify({"status": "error", "message": "无权操作"}), 403
    
    data = request.json
    doc_id = data.get("doc_id")
    if not doc_id:
        return jsonify({"status": "error", "message": "缺少文章ID"}), 400
    
    articles = cdata.get("articles", [])
    if doc_id in articles:
        return jsonify({"status": "error", "message": "文章已在合集中"}), 400
    
    articles.append(doc_id)
    cdata["articles"] = articles
    save_collection(collection_id, cdata)
    return jsonify({"status": "ok"})

@app.route("/api/collections/<collection_id>/articles/<doc_id>", methods=["DELETE"])
@login_required
def remove_article_from_collection(collection_id, doc_id):
    collections = load_collections()
    cdata = collections.get(collection_id)
    if not cdata:
        return jsonify({"status": "error", "message": "合集不存在"}), 404
    if not can_edit_collection(cdata, g.username, g.role):
        return jsonify({"status": "error", "message": "无权操作"}), 403
    
    articles = cdata.get("articles", [])
    if doc_id in articles:
        articles.remove(doc_id)
    cdata["articles"] = articles
    save_collection(collection_id, cdata)
    return jsonify({"status": "ok"})

@app.route("/api/collections/<collection_id>/editors", methods=["POST"])
@login_required
def add_collection_editor(collection_id):
    """添加协作者（仅作者/owner 可操作）"""
    collections = load_collections()
    cdata = collections.get(collection_id)
    if not cdata:
        return jsonify({"status": "error", "message": "合集不存在"}), 404
    if not can_edit_collection(cdata, g.username, g.role):
        return jsonify({"status": "error", "message": "无权操作"}), 403

    data = request.json

    query = data.get("username", "").strip()
    if not query:
        return jsonify({"status": "error", "message": "缺少用户名或笔名"}), 400

    users = load_users()
    # 先按用户名查，再按笔名查
    username = query if query in users else None
    if not username:
        for uname, uinfo in users.items():
            nick = uinfo.get("nickname", "") or ""
            if str(nick).strip() == query:
                username = uname
                break
    if not username:
        return jsonify({"status": "error", "message": f"用户不存在：{query}"}), 404

    editors = cdata.get("editors", [])
    if username in editors:
        return jsonify({"status": "error", "message": "该用户已是协作者"}), 400

    editors.append(username)
    cdata["editors"] = editors
    save_collection(collection_id, cdata)
    return jsonify({"status": "ok", "editors": editors})


@app.route("/api/collections/<collection_id>/editors/<editor_name>", methods=["DELETE"])
@login_required
def remove_collection_editor(collection_id, editor_name):
    """移除协作者（仅作者/owner 可操作）"""
    collections = load_collections()
    cdata = collections.get(collection_id)
    if not cdata:
        return jsonify({"status": "error", "message": "合集不存在"}), 404
    if not can_edit_collection(cdata, g.username, g.role):
        return jsonify({"status": "error", "message": "无权操作"}), 403

    editors = cdata.get("editors", [])
    if editor_name in editors:
        editors.remove(editor_name)
    cdata["editors"] = editors
    save_collection(collection_id, cdata)
    return jsonify({"status": "ok", "editors": editors})


@app.route("/api/my_collections", methods=["GET"])
@login_required
def my_collections():
    collections = load_collections()
    result = []
    for cid, cdata in collections.items():
        if cdata.get("author") == g.username:
            result.append({
                "id": cid,
                "name": cdata.get("name", "未命名合集"),
                "description": cdata.get("description", ""),
                "created_at": cdata.get("created_at", ""),
                "article_count": len(cdata.get("articles", []))
            })
    result.sort(key=lambda x: x["created_at"], reverse=True)
    return jsonify(result)

# ---------- 管理员 ----------
@app.route("/api/admin/set_admin", methods=["POST"])
@login_required
def set_admin():
    if g.role != "owner":
        return jsonify({"status": "error", "message": "仅站长可操作"}), 403
    data = request.json
    target = data.get("username")
    action = data.get("action")
    if not target or action not in ("promote", "demote"):
        return jsonify({"status": "error", "message": "参数错误"}), 400
    users = load_users()
    if target not in users:
        return jsonify({"status": "error", "message": "用户不存在"}), 404
    if action == "promote":
        users[target]["role"] = "admin"
    else:
        users[target]["role"] = "user"
    save_users(users)
    return jsonify({"status": "ok", "message": f"已{'设为' if action=='promote' else '取消'}管理员"})

@app.route("/api/admin/set_dojin", methods=["POST"])
@login_required
def set_dojin():
    if g.role != "owner":
        return jsonify({"status": "error", "message": "仅站长可操作"}), 403
    data = request.json
    target = data.get("username")
    action = data.get("action")  # "set" or "unset"
    if not target or action not in ("set", "unset"):
        return jsonify({"status": "error", "message": "参数错误"}), 400
    users = load_users()
    if target not in users:
        return jsonify({"status": "error", "message": "用户不存在"}), 404
    if action == "set":
        users[target]["dojin"] = True
        msg = "已设为同人"
    else:
        users[target].pop("dojin", None)
        msg = "已取消同人"
    save_users(users)
    return jsonify({"status": "ok", "message": msg})

# ---------- 文档 API ----------
@app.route("/api/docs")
def list_docs():
    sort_by = request.args.get("sort", "updated")
    cache_key = f'docs_list_{sort_by}'
    cached = cache_get(cache_key, max_age=10)
    if cached is not None:
        return jsonify(cached)

    docs = []
    try:
        entries = os.listdir(DOCS_DIR)
    except OSError:
        return jsonify([])

    users = load_users()
    # 建立 nickname -> avatar 映射
    _nickname_avatar = {}
    for u, info in users.items():
        nick = info.get("nickname", u)
        av = info.get("avatar", "")
        if av and av != "None" and av != "null":
            _nickname_avatar[nick] = av
    for filename in entries:
        if not filename.endswith(".txt") or filename.startswith("."):
            continue
        doc_id = filename.replace(".txt", "")
        meta = load_meta(doc_id)
        filepath = os.path.join(DOCS_DIR, filename)
        try:
            updated = datetime.fromtimestamp(os.path.getmtime(filepath)).isoformat()
        except OSError:
            updated = datetime.now().isoformat()
        # 获取作者头像和角色（通过 nickname 映射）
        authors = meta.get("authors", [])
        author_avatars = {}
        author_roles = {}
        for a in authors:
            if a in _nickname_avatar:
                author_avatars[a] = _nickname_avatar[a]
            # 找作者的 role 和 dojin 状态
            for u, info in users.items():
                if info.get("nickname") == a:
                    role = info.get("role", "user")
                    if role != "user":
                        author_roles[a] = role
                    if info.get("dojin"):
                        author_roles[a] = author_roles.get(a, "") + "_dojin"
                    break
        docs.append({
            "text_stroke": meta.get("text_stroke", False),
            "cover": meta.get("cover", ""),
            "cover_thumb": meta.get("cover_thumb", ""),
            "title_color": meta.get("title_color", "#000000"),
            "prompt_color": meta.get("prompt_color", "#0f0b0b"),
            "id": doc_id,
            "number": meta.get("number", 999999),
            "views": meta.get("views", 0),
            "authors": authors,
            "author_avatars": author_avatars,
            "author_roles": author_roles,
            "title": meta.get("title", doc_id),
            "status": meta.get("status", "draft"),
            "updated": updated
        })
    if sort_by == "views":
        docs.sort(key=lambda x: x["views"], reverse=True)
    elif sort_by == "number":
        docs.sort(key=lambda x: x["number"], reverse=False)
    else:
        docs.sort(key=lambda x: x["updated"], reverse=True)
    cache_set(cache_key, docs)
    return jsonify(docs)

@app.route("/api/my_drafts")
@login_required
def my_drafts():
    username = g.username
    role = g.role
    users = load_users()
    nickname = users.get(username, {}).get("nickname", username)
    
    drafts = []
    for filename in os.listdir(DOCS_DIR):
        if filename.endswith(".txt") and not filename.startswith("."):
            doc_id = filename.replace(".txt", "")
            meta = load_meta(doc_id)
            if meta.get("status") != "draft":
                continue
            if role != "owner" and nickname not in meta.get("authors", []):
                continue
            drafts.append({
                "id": doc_id,
                "number": meta.get("number", 999999),
                "title": meta.get("title", doc_id),
                "authors": meta.get("authors", []),
                "views": meta.get("views", 0),
                "cover": meta.get("cover", ""),
                "cover_thumb": meta.get("cover_thumb", ""),
                "prompt_color": meta.get("prompt_color", "#0f0b0b"),
                "title_color": meta.get("title_color", "#000000"),
                "text_stroke": meta.get("text_stroke", False),
                "status": "draft"
            })
    drafts.sort(key=lambda x: x["number"])
    return jsonify(drafts)

@app.route("/api/search")
def search_docs():
    get_optional_user()
    q = request.args.get("q", "").strip()
    if not q:
        return jsonify([])
    users = load_users()
    # 建立 nickname -> avatar 映射
    _nickname_avatar = {}
    for u, info in users.items():
        nick = info.get("nickname", u)
        av = info.get("avatar", "")
        if av and av != "None" and av != "null":
            _nickname_avatar[nick] = av
    results = []
    for filename in os.listdir(DOCS_DIR):
        if filename.endswith(".txt") and not filename.startswith("."):
            doc_id = filename.replace(".txt", "")
            meta = load_meta(doc_id)
            status = meta.get("status", "published")
            if status == "draft":
                if not g.username:
                    continue
                user_nickname = users.get(g.username, {}).get("nickname", g.username)
                if user_nickname not in meta.get("authors", []):
                    continue
            title = meta.get("title", doc_id)
            path = os.path.join(DOCS_DIR, f"{doc_id}.txt")
            content = ""
            if os.path.exists(path):
                with open(path, "r", encoding="utf-8") as f:
                    content = f.read()
            if (q.lower() in title.lower() or
                q.lower() in content.lower() or
                (meta.get("number") and str(meta["number"]) == q)):
                authors = meta.get("authors", [])
                author_avatars = {}
                author_roles = {}
                for a in authors:
                    if a in _nickname_avatar:
                        author_avatars[a] = _nickname_avatar[a]
                    for u, info in users.items():
                        if info.get("nickname") == a:
                            role = info.get("role", "user")
                            if role != "user":
                                author_roles[a] = role
                            if info.get("dojin"):
                                author_roles[a] = author_roles.get(a, "") + "_dojin"
                            break
                results.append({
                    "id": doc_id,
                    "number": meta.get("number", 0),
                    "title": title,
                    "authors": authors,
                    "author_avatars": author_avatars,
                    "author_roles": author_roles,
                    "views": meta.get("views", 0),
                    "cover": meta.get("cover", ""),
                    "cover_thumb": meta.get("cover_thumb", ""),
                    "prompt_color": meta.get("prompt_color", "#0f0b0b"),
                    "title_color": meta.get("title_color", "#000000"),
                    "text_stroke": meta.get("text_stroke", False)
                })
    results.sort(key=lambda x: x["number"])
    return jsonify(results)

@app.route("/api/search/users")
def search_users():
    q = request.args.get("q", "").strip()
    if not q:
        return jsonify([])
    users = load_users()
    results = []
    for username, info in users.items():
        if q.lower() in username.lower() or q.lower() in info.get("nickname", "").lower():
            results.append({
                "username": username,
                "nickname": info.get("nickname", username),
                "avatar": info.get("avatar"),
                "role": info.get("role", "user")
            })
    return jsonify(results)

@app.route("/api/doc/<doc_id>")
def get_doc(doc_id):
    get_optional_user()
    path = os.path.join(DOCS_DIR, f"{doc_id}.txt")
    text_content = ""
    if os.path.exists(path):
        with open(path, "r", encoding="utf-8") as f:
            text_content = f.read()
    meta = load_meta(doc_id)
    status = meta.get("status", "published")
    users = load_users()
    username = g.username
    if status == "draft":
        if not username:
            # 未登录用户，检查备份
            backup = published_backup_path(doc_id)
            if os.path.exists(backup):
                with open(backup, "r", encoding="utf-8") as f:
                    text_content = f.read()
                # 不增加阅读量
                return jsonify({
                    "content": text_content,
                    "title": meta.get("title", doc_id),
                    "number": meta.get("number"),
                    "authors": meta.get("authors", []),
                    "status": "draft",
                    "is_published_version": True,   # 标记是旧版
                    # 其他必要字段可补上
                })
            return jsonify({"status": "error", "message": "作者正在奋笔疾书..."}), 404
        # 已登录用户，判断权限...
        user_info = users.get(username, {})
        user_nickname = user_info.get("nickname", username)
        user_is_author = (user_nickname in meta.get("authors", [])) or (username in meta.get("authors", []))
        if not user_is_author and g.role not in ("admin", "owner"):
            # 也无权限，返回备份
            backup = published_backup_path(doc_id)
            if os.path.exists(backup):
                with open(backup, "r", encoding="utf-8") as f:
                    text_content = f.read()
                return jsonify({
                    "content": text_content,
                    "title": meta.get("title", doc_id),
                    "number": meta.get("number"),
                    "authors": meta.get("authors", []),
                    "status": "draft",
                    "is_published_version": True,
                })
            return jsonify({"status": "error", "message": "作者正在奋笔疾书..."}), 404
    meta["views"] = meta.get("views", 0) + 1
    save_meta(doc_id, meta)
    author_usernames = []
    for nick in meta.get("authors", []):
        for u, info in users.items():
            if info.get("nickname") == nick:
                author_usernames.append(u)
                break
    # 查找下一章（遍历该文章所属的所有合集）
    next_in_collection = None
    next_collections = []
    collections = load_collections()
    for cid, cdata in collections.items():
        articles = cdata.get("articles", [])
        if doc_id in articles:
            idx = articles.index(doc_id)
            if idx + 1 < len(articles):
                next_id = articles[idx + 1]
                next_meta = load_meta(next_id)
                next_collections.append({
                    "id": next_id,
                    "title": next_meta.get("title", next_id),
                    "collection_name": cdata.get("name", cid),
                    "collection_id": cid
                })
    # 只有一个时保持单值，多个时返回数组
    if len(next_collections) == 1:
        next_in_collection = next_collections[0]
    elif len(next_collections) > 1:
        next_in_collection = next_collections
    return jsonify({
        "text_stroke": meta.get("text_stroke", False),
        "cover": meta.get("cover", ""),
        "cover_thumb": meta.get("cover_thumb", ""),
        "title_color": meta.get("title_color", "#000000"),
        "prompt_color": meta.get("prompt_color", "#0f0b0b"),
        "content": text_content,
        "authors": meta.get("authors", []),
        "number": meta.get("number"),
        "views": meta.get("views", 0),
        "likes": meta.get("likes", 0),
        "liked_by": meta.get("liked_by", []),
        "edit_mode": meta.get("edit_mode", "public"),
        "allowed_users": meta.get("allowed_users", []),
        "title": meta.get("title", doc_id),
        "author_usernames": author_usernames,
        "status": status,
        "role": users.get(username, {}).get("role", "user") if username else "user",
        "next_in_collection": next_in_collection
    })
@app.route("/api/doc/<doc_id>/like", methods=["POST"])
@login_required
def toggle_like(doc_id):
    username = g.username
    meta = load_meta(doc_id)
    if "liked_by" not in meta:
        meta["liked_by"] = []
    if "likes" not in meta:
        meta["likes"] = 0
    if username in meta["liked_by"]:
        meta["liked_by"].remove(username)
        meta["likes"] = max(0, meta["likes"] - 1)
        liked = False
    else:
        meta["liked_by"].append(username)
        meta["likes"] = meta["likes"] + 1
        liked = True
    save_meta(doc_id, meta)
    return jsonify({"status": "ok", "likes": meta["likes"], "liked": liked})

def published_backup_path(doc_id):
    return os.path.join(DOCS_DIR, f"{doc_id}.published.bak")   # 改后缀

@app.route("/api/doc/<doc_id>", methods=["POST"])
@login_required
def save_doc(doc_id):
    username = g.username
    if not check_edit_permission(doc_id, username):
        return jsonify({"status": "error", "message": "无权编辑此文章"}), 403

    path = os.path.join(DOCS_DIR, f"{doc_id}.txt")

    # 🔥 先备份旧内容（如果文章是发布状态且即将变为草稿）
    meta = load_meta(doc_id)
    old_status = meta.get("status", "published")
    new_status = request.json.get("status")

    if old_status == "published" and new_status == "draft":
        if os.path.exists(path):
            backup_path = published_backup_path(doc_id)
            with open(path, "r", encoding="utf-8") as src:
                content_to_backup = src.read()
            with open(backup_path, "w", encoding="utf-8") as dst:
                dst.write(content_to_backup)

    # 现在写入新内容
    with open(path, "w", encoding="utf-8") as f:
        f.write(request.json.get("content", ""))

    # 更新状态
    if new_status in ("draft", "published"):
        if new_status == "published":
            backup_path = published_backup_path(doc_id)
            if os.path.exists(backup_path):
                os.remove(backup_path)
        meta["status"] = new_status
        save_meta(doc_id, meta)

    return jsonify({"status": "ok"})

@app.route("/api/doc/<doc_id>", methods=["DELETE"])
@login_required
def delete_doc(doc_id):
    username = g.username
    meta = load_meta(doc_id)
    users = load_users()
    nickname = users.get(username, {}).get("nickname", username)
    role = g.role
    authors = meta.get("authors", [])
    backup = published_backup_path(doc_id)
    if os.path.exists(backup):
        os.remove(backup)
    if role == "owner":
        pass
    elif role == "admin" and len(authors) == 0:
        pass
    else:
        if nickname not in authors:
            return jsonify({"status": "error", "message": "只有作者可以删除"}), 403
    path = os.path.join(DOCS_DIR, f"{doc_id}.txt")
    meta_path = get_meta_path(doc_id)
    comments_path = get_comments_path(doc_id)
    # 清理封面文件
    _cleanup_old_cover(meta)
    for p in [path, meta_path, comments_path]:
        if os.path.exists(p):
            os.remove(p)
    # 从所有合集中移除
    collections = load_collections()
    changed = False
    for cid, cdata in collections.items():
        articles = cdata.get("articles", [])
        if doc_id in articles:
            articles.remove(doc_id)
            cdata["articles"] = articles
            save_collection(cid, cdata)
            changed = True

    # 清理临时缓存
    cache_del('docs_list_updated')
    cache_del('docs_list_number')
    cache_del('docs_list_views')
    return jsonify({"status": "ok"})

def _cleanup_old_cover(meta):
    """删除旧封面文件"""
    for key in ('cover', 'cover_thumb'):
        path = meta.get(key, '')
        if path:
            filepath = os.path.join('static', path.lstrip('/'))
            if os.path.exists(filepath):
                try:
                    os.remove(filepath)
                except OSError:
                    pass

@app.route("/api/doc/<doc_id>/cover", methods=["POST"])
@login_required
def upload_cover(doc_id):
    username = g.username
    meta = load_meta(doc_id)
    users = load_users()
    nickname = users.get(username, {}).get("nickname", username)
    if nickname not in meta.get("authors", []) and g.role not in ("admin", "owner"):
        return jsonify({"status": "error", "message": "无权操作"}), 403

    if 'cover' not in request.files:
        return jsonify({"status": "error", "message": "没有文件"}), 400
    file = request.files['cover']
    if file.filename == '':
        return jsonify({"status": "error", "message": "空文件名"}), 400

    # 大小检查
    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    if size > UPLOAD_MAX_SIZE:
        return jsonify({"status": "error", "message": f"文件过大，最大允许{UPLOAD_MAX_SIZE//1024//1024}MB"}), 400

    ext = file.filename.rsplit('.', 1)[-1].lower()
    if ext not in allowed_ext():
        return jsonify({"status": "error", "message": "不支持的格式"}), 400

    # 清理旧封面
    _cleanup_old_cover(meta)

    # 保存原图（并压缩）
    filename = f"{doc_id}_{uuid.uuid4().hex[:8]}.{ext}"
    filepath = os.path.join(COVERS_DIR, filename)
    file.save(filepath)
    optimize_image(filepath)

    # 生成缩略图
    thumb_filename = f"{doc_id}_{uuid.uuid4().hex[:8]}_thumb.webp"
    thumb_filepath = os.path.join(COVERS_DIR, thumb_filename)
    make_thumb(filepath, thumb_filepath, max_size=300)

    # 更新 meta
    meta['cover'] = f"/static/covers/{filename}"
    meta['cover_thumb'] = f"/static/covers/{thumb_filename}"
    save_meta(doc_id, meta)

    return jsonify({"status": "ok", "cover": meta['cover']})

@app.route("/api/doc/<doc_id>/colors", methods=["POST"])
@login_required
def set_colors(doc_id):
    username = g.username
    meta = load_meta(doc_id)
    users = load_users()
    nickname = users.get(username, {}).get("nickname", username)
    if nickname not in meta.get("authors", []) and g.role not in ("admin", "owner"):
        return jsonify({"status": "error", "message": "无权操作"}), 403
    data = request.json
    meta["prompt_color"] = data.get("prompt_color", "#0f0b0b")
    meta["title_color"] = data.get("title_color", "#000000")
    save_meta(doc_id, meta)
    return jsonify({"status": "ok"})

@app.route("/api/doc/<doc_id>/text_stroke", methods=["POST"])
@login_required
def set_text_stroke(doc_id):
    username = g.username
    meta = load_meta(doc_id)
    users = load_users()
    nickname = users.get(username, {}).get("nickname", username)
    if nickname not in meta.get("authors", []) and g.role not in ("admin", "owner"):
        return jsonify({"status": "error", "message": "无权操作"}), 403
    enabled = request.json.get("enabled", False)
    meta["text_stroke"] = enabled
    save_meta(doc_id, meta)
    return jsonify({"status": "ok", "text_stroke": enabled})

@app.route("/api/doc/<doc_id>/permission", methods=["POST"])
@login_required
def update_permission(doc_id):
    username = g.username
    meta = load_meta(doc_id)
    users = load_users()
    nickname = users.get(username, {}).get("nickname", username)
    if g.role not in ("admin", "owner"):
        if nickname != (meta.get("authors", [])[0] if meta.get("authors") else None):
            return jsonify({"status": "error", "message": "只有第一作者可以修改权限"}), 403
    data = request.json
    mode = data.get("mode")
    allowed = data.get("allowed_users", [])
    if mode not in ("public", "authors", "restricted"):
        return jsonify({"status": "error", "message": "无效的权限模式"}), 400
    meta["edit_mode"] = mode
    meta["allowed_users"] = allowed
    save_meta(doc_id, meta)
    return jsonify({"status": "ok"})

@app.route("/api/doc/<doc_id>/title", methods=["POST"])
@login_required
def update_title(doc_id):
    username = g.username
    meta = load_meta(doc_id)
    users = load_users()
    nickname = users.get(username, {}).get("nickname", username)
    if g.role not in ("admin", "owner"):
        if nickname != (meta.get("authors", [])[0] if meta.get("authors") else None):
            return jsonify({"status": "error", "message": "只有第一作者可以修改标题"}), 403
    new_title = request.json.get("title", "").strip()
    if not new_title:
        return jsonify({"status": "error", "message": "标题不能为空"}), 400
    meta["title"] = new_title
    save_meta(doc_id, meta)
    return jsonify({"status": "ok", "title": new_title})

@app.route("/api/create_doc", methods=["POST"])
@login_required
def create_doc():
    username = g.username
    number = get_next_number()
    doc_id = str(number)
    users = load_users()
    user_info = users.get(username, {})
    author = user_info.get("nickname", username)
    content_path = os.path.join(DOCS_DIR, f"{doc_id}.txt")
    if not os.path.exists(content_path):
        with open(content_path, "w", encoding="utf-8") as f:
            f.write("")
    meta = load_meta(doc_id)
    meta["number"] = number
    meta["authors"] = [author]
    meta["views"] = meta.get("views", 0)
    meta["edit_mode"] = meta.get("edit_mode", "public")
    meta["allowed_users"] = meta.get("allowed_users", [])
    meta["title"] = f"第{number}号文章"
    meta["status"] = "draft"
    save_meta(doc_id, meta)
    return jsonify({"status": "ok", "number": number, "doc_id": doc_id})

# ---------- 文档下载 ----------
@app.route("/api/doc/<doc_id>/download/<fmt>")
def download_doc(doc_id, fmt):
    """下载文章为 txt / docx 格式"""
    if fmt not in ('txt', 'docx'):
        return jsonify({"status": "error", "message": "不支持的格式"}), 400

    path = os.path.join(DOCS_DIR, f"{doc_id}.txt")
    if not os.path.exists(path):
        return jsonify({"status": "error", "message": "文章不存在"}), 404

    with open(path, 'r', encoding='utf-8') as f:
        text_content = f.read()

    meta = load_meta(doc_id)
    title = meta.get('title', f'第{doc_id}号文章')
    authors = meta.get('authors', [])
    number = meta.get('number')

    if fmt == 'txt':
        header = f"{title}\n"
        if number:
            header = f"#{number} {title}\n"
        if authors:
            header += f"作者：{'、'.join(authors)}\n"
        header += "=" * 40 + "\n\n"
        full = header + text_content
        from io import BytesIO
        buf = BytesIO()
        buf.write(full.encode('utf-8'))
        buf.seek(0)
        safe_title = re.sub(r'[\\/:*?"<>|]', '_', title)
        return send_file(buf, as_attachment=True, download_name=f"{safe_title}.txt", mimetype='text/plain; charset=utf-8')

    elif fmt == 'docx':
        from docx import Document
        from docx.shared import Pt
        doc = Document()
        # 标题
        heading_text = title
        if number:
            heading_text = f"#{number} {title}"
        doc.add_heading(heading_text, level=1)
        # 作者
        if authors:
            doc.add_paragraph(f"作者：{'、'.join(authors)}").style = doc.styles['Normal']
        # 分割线
        doc.add_paragraph("_" * 40)
        # 正文
        for para in text_content.split('\n'):
            if para.strip():
                doc.add_paragraph(para)
            else:
                doc.add_paragraph('')
        from io import BytesIO
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        safe_title = re.sub(r'[\\/:*?"<>|]', '_', title)
        return send_file(buf, as_attachment=True, download_name=f"{safe_title}.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

# ---------- 文档上传 ----------
@app.route("/api/upload_docfile", methods=["POST"])
@login_required
def upload_docfile():
    """上传 .txt / .doc / .docx 文档，提取文本内容返回"""
    username = g.username
    if 'docfile' not in request.files:
        return jsonify({"status": "error", "message": "没有文件"}), 400
    file = request.files['docfile']
    if file.filename == '':
        return jsonify({"status": "error", "message": "空文件名"}), 400

    ext = file.filename.rsplit('.', 1)[-1].lower()
    if ext not in ('txt', 'doc', 'docx'):
        return jsonify({"status": "error", "message": "仅支持 .txt .doc .docx 格式"}), 400

    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    if size > 20 * 1024 * 1024:
        return jsonify({"status": "error", "message": "文件过大，最大允许 20MB"}), 400

    try:
        if ext == 'txt':
            content = file.read().decode('utf-8', errors='replace')
        elif ext == 'docx':
            import zipfile, xml.etree.ElementTree as ET
            zf = zipfile.ZipFile(file)
            xml_content = zf.read('word/document.xml')
            root = ET.fromstring(xml_content)
            ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            paragraphs = []
            for p in root.iter(f'{{{ns}}}p'):
                texts = []
                for t in p.iter(f'{{{ns}}}t'):
                    if t.text:
                        texts.append(t.text)
                paragraphs.append(''.join(texts))
            content = '\n'.join(paragraphs)
        elif ext == 'doc':
            raw = file.read()
            readable = []
            buf = []
            for b in raw:
                if 32 <= b < 127 or b in (10, 13) or b >= 128:
                    buf.append(chr(b))
                else:
                    if buf:
                        readable.append(''.join(buf))
                        buf = []
            if buf:
                readable.append(''.join(buf))
            content = ' '.join(readable)
            content = re.sub(r'\s+', ' ', content).strip()
        else:
            return jsonify({"status": "error", "message": "不支持的格式"}), 400

        if not content.strip():
            return jsonify({"status": "error", "message": "未能从文件中提取到文本内容"}), 400

        return jsonify({"status": "ok", "content": content, "filename": file.filename})
    except Exception as e:
        return jsonify({"status": "error", "message": f"文件解析失败：{str(e)}"}), 400

# ---------- 图片上传 ----------
@app.route("/api/upload_image", methods=["POST"])
@login_required
def upload_image():
    username = g.username
    if 'image' not in request.files:
        return jsonify({"status": "error", "message": "没有文件"}), 400
    file = request.files['image']
    if file.filename == '':
        return jsonify({"status": "error", "message": "空文件名"}), 400

    # 大小检查
    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    if size > UPLOAD_MAX_SIZE:
        return jsonify({"status": "error", "message": f"文件过大，最大允许{UPLOAD_MAX_SIZE//1024//1024}MB"}), 400

    ext = file.filename.rsplit('.', 1)[-1].lower()
    if ext not in allowed_ext():
        return jsonify({"status": "error", "message": "不支持的格式"}), 400

    filename = f"{uuid.uuid4().hex[:12]}.{ext}"
    filepath = os.path.join(UPLOADS_DIR, filename)
    file.save(filepath)
    optimize_image(filepath)
    return jsonify({"status": "ok", "url": f"/static/uploads/{filename}"})

@app.route("/api/upload_video", methods=["POST"])
@login_required
def upload_video():
    username = g.username
    if 'video' not in request.files:
        return jsonify({"status": "error", "message": "没有文件"}), 400
    file = request.files['video']
    if file.filename == '':
        return jsonify({"status": "error", "message": "空文件名"}), 400

    # 大小检查（100MB）
    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    if size > 100 * 1024 * 1024:
        return jsonify({"status": "error", "message": "视频文件过大，最大允许100MB"}), 400

    ext = file.filename.rsplit('.', 1)[-1].lower()
    if ext not in allowed_video_ext():
        return jsonify({"status": "error", "message": "不支持的视频格式，支持 mp4, webm, mov, avi"}), 400

    # 先保存临时文件
    raw_filename = f"{uuid.uuid4().hex[:12]}.{ext}"
    raw_path = os.path.join(VIDEOS_DIR, raw_filename)
    file.save(raw_path)

    # 用 ffmpeg 转码为 H.264 兼容格式
    out_filename = f"{uuid.uuid4().hex[:12]}.mp4"
    out_path = os.path.join(VIDEOS_DIR, out_filename)
    try:
        import subprocess
        result = subprocess.run([
            "ffmpeg", "-i", raw_path,
            "-c:v", "libx264", "-preset", "fast", "-crf", "23",
            "-c:a", "aac", "-b:a", "128k",
            "-movflags", "+faststart",
            "-y", out_path
        ], capture_output=True, text=True, timeout=120)
        if result.returncode == 0:
            os.remove(raw_path)  # 删除原始文件
            return jsonify({"status": "ok", "url": f"/static/uploads/videos/{out_filename}"})
        else:
            # 转码失败，回退用原始文件
            os.rename(raw_path, out_path)
            return jsonify({"status": "ok", "url": f"/static/uploads/videos/{out_filename}", "note": "转码失败，原始文件已保留"})
    except Exception:
        # ffmpeg 超时或其他错误，回退
        os.rename(raw_path, out_path)
        return jsonify({"status": "ok", "url": f"/static/uploads/videos/{out_filename}", "note": "转码失败，原始文件已保留"})

# ---------- 错字检查 ----------
DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY", "")
DEEPSEEK_URL = "https://api.deepseek.com/v1/chat/completions"

# 加载本地错别字字典
_SPELL_DICT = None

def load_spell_dict():
    """加载本地错别字字典，返回 {错字: 正确字}"""
    d = {}
    dict_path = os.path.join(os.path.dirname(__file__), "spell_dict.csv")
    if not os.path.exists(dict_path):
        return d
    with open(dict_path, "r", encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            parts = line.split("\t")
            if len(parts) == 2:
                wrong, correct = parts[0].strip(), parts[1].strip()
                if wrong and correct:
                    d[wrong] = correct
    return d

def local_spellcheck(text):
    """本地字典扫描，返回错误列表"""
    global _SPELL_DICT
    if _SPELL_DICT is None:
        _SPELL_DICT = load_spell_dict()
    errors = []
    # 按错误长度降序匹配，避免短词覆盖长词
    for wrong, correct in sorted(_SPELL_DICT.items(), key=lambda x: -len(x[0])):
        start = 0
        while True:
            idx = text.find(wrong, start)
            if idx == -1:
                break
            errors.append({
                "word": wrong,
                "start": idx,
                "end": idx + len(wrong),
                "suggestion": correct
            })
            start = idx + 1
    return errors

def merge_errors(local_errs, ai_errs):
    """合并本地和AI结果，去重（按 (start,end) 去重）"""
    seen_spans = set()
    merged = []
    for e in local_errs + ai_errs:
        key = (e["start"], e["end"])
        if key not in seen_spans:
            seen_spans.add(key)
            merged.append(e)
    merged.sort(key=lambda x: x["start"])
    return merged

@app.route("/api/spellcheck", methods=["POST"])
@login_required
def api_spellcheck():
    # 白名单：这些词是正确的常用中文词汇，AI 如果报了就是幻觉
    WHITELIST = {
        "知道", "管道", "道理", "到了", "再来", "再见", "在家",
        "在吗", "在手", "在场", "在外", "在内", "在线",
        "在此", "在我", "何在", "存在", "自在", "现在",
        "正在", "内在", "潜在", "实在", "好在", "在于",
        "告诉", "报告", "被告", "通知", "公告", "告知",
        "速报", "报道", "预报", "举报", "海报", "申报",
        "谎报", "到来", "达到", "到底", "到达", "等到",
    }

    data = request.get_json()
    text = (data.get("text") or "").strip()
    if not text or len(text) > 10000:
        return jsonify({"status": "error", "message": "文本为空或过长"}), 400

    # 第一步：本地字典快速扫描（毫秒级）
    local_errors = local_spellcheck(text)

    # 第二步：AI 深度检查
    ai_errors = []  # 暂时禁用 AI 层，只使用本地字典（避免幻觉）



    # 第三步：合并去重（本地结果优先，AI 结果作补充）
    merged = merge_errors(local_errors, ai_errors)

    return jsonify({"status": "ok", "errors": merged, "count": len(merged)})

# ---------- PWA Service Worker ----------
@app.route("/sw.js")
def service_worker():
    response = make_response(send_from_directory("static", "sw.js"))
    response.headers["Content-Type"] = "application/javascript"
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    response.headers["Service-Worker-Allowed"] = "/"
    return response

# ---------- 图片缓存 ----------
@app.after_request
def add_cache_header(response):
    path = request.path
    if path.startswith('/static/covers') or path.startswith('/avatars') or path.startswith('/static/uploads/videos') or path.startswith('/static/uploads'):
        response.headers['Cache-Control'] = 'public, max-age=604800'
    return response

# ---------- 用户与关注 ----------
@app.route("/api/user/<path:username>")
def api_user_profile(username):
    get_optional_user()
    users = load_users()
    if username not in users:
        return jsonify({"status": "error", "message": "用户不存在"}), 404
    user = users[username]
    docs = []
    for filename in os.listdir(DOCS_DIR):
        if filename.endswith(".txt") and not filename.startswith("."):
            doc_id = filename.replace(".txt", "")
            meta = load_meta(doc_id)
            if user["nickname"] in meta.get("authors", []):
                docs.append({
                    "id": doc_id,
                    "number": meta.get("number", 0),
                    "title": meta.get("title", doc_id),
                    "cover": meta.get("cover", ""),
                    "cover_thumb": meta.get("cover_thumb", ""),
                    "authors": meta.get("authors", [])
                })
    current_user_name = g.username
    is_following = False
    if current_user_name and current_user_name in users:
        current_following = users[current_user_name].get("following", [])
        is_following = username in current_following
    return jsonify({
        "username": username,
        "nickname": user.get("nickname", username),
        "avatar": user.get("avatar"),
        "role": user.get("role", "user"),
        "is_dojin": user.get("dojin", False),
        "articles": docs,
        "followers_count": len(user.get("followers", [])),
        "following_count": len(user.get("following", [])),
        "is_following": is_following
    })

@app.route("/api/user/<path:username>/following")
def get_following(username):
    users = load_users()
    if username not in users:
        return jsonify([])
    target = users[username]
    result = []
    for u in target.get("following", []):
        info = users.get(u)
        if info:
            result.append({"username": u, "nickname": info.get("nickname", u), "avatar": info.get("avatar")})
    return jsonify(result)

@app.route("/api/admin/user_list")
@login_required
def api_admin_user_list():
    if g.role != "owner":
        return jsonify({"status": "error", "message": "仅站长可操作"}), 403
    users = load_users()
    result = []
    for u, info in users.items():
        result.append({
            "username": u,
            "nickname": info.get("nickname", u),
            "role": info.get("role", "user"),
            "is_dojin": info.get("dojin", False)
        })
    return jsonify({"users": result})

@app.route("/api/user/<path:username>/followers")
def get_followers(username):
    users = load_users()
    if username not in users:
        return jsonify([])
    target = users[username]
    result = []
    for u in target.get("followers", []):
        info = users.get(u)
        if info:
            result.append({"username": u, "nickname": info.get("nickname", u), "avatar": info.get("avatar")})
    return jsonify(result)

@app.route("/api/follow/<path:username>", methods=["POST"])
@login_required
def toggle_follow(username):
    current_username = g.username
    users = load_users()
    if current_username not in users or username not in users:
        return jsonify({"status": "error", "message": "用户不存在"}), 404
    if current_username == username:
        return jsonify({"status": "error", "message": "不能关注自己"}), 400
    current_following = users[current_username].setdefault("following", [])
    target_followers = users[username].setdefault("followers", [])
    if username in current_following:
        current_following.remove(username)
        target_followers.remove(current_username)
        action = "unfollowed"
    else:
        current_following.append(username)
        target_followers.append(current_username)
        action = "followed"
    save_users(users)
    return jsonify({"status": "ok", "action": action, "followers_count": len(target_followers)})

@app.route("/api/mutual_friends")
@login_required
def mutual_friends():
    username = g.username
    users = load_users()
    if username not in users:
        return jsonify([])
    following = set(users[username].get("following", []))
    followers = set(users[username].get("followers", []))
    mutual = following & followers
    result = []
    for u in mutual:
        info = users.get(u)
        if info:
            result.append({
                "username": u,
                "nickname": info.get("nickname", u),
                "avatar": info.get("avatar")
            })
    return jsonify(result)

# ---------- 通知 ----------
@app.route("/api/notifications", methods=["GET"])
@login_required
def get_notifications():
    username = g.username
    users = load_users()
    return jsonify(users.get(username, {}).get('notifications', []))

@app.route("/api/notifications/read", methods=["POST"])
@login_required
def mark_notification_read():
    username = g.username
    data = request.get_json(silent=True) or {}
    noti_id = data.get('notification_id')
    users = load_users()
    if username not in users:
        return jsonify({"status": "error"}), 404
    notifications = users[username].get('notifications', [])
    if noti_id:
        for n in notifications:
            if n['id'] == noti_id:
                n['read'] = True
                break
    else:
        for n in notifications:
            n['read'] = True
    save_users(users)
    return jsonify({"status": "ok"})

@app.route("/api/notifications/read_all", methods=["POST"])
@login_required
def mark_all_notifications_read():
    username = g.username
    users = load_users()
    if username in users:
        for n in users[username].get('notifications', []):
            n['read'] = True
        save_users(users)
    return jsonify({"status": "ok"})

@app.route("/api/notifications/count", methods=["GET"])
@login_required
def unread_count():
    username = g.username
    users = load_users()
    unread = sum(1 for n in users.get(username, {}).get('notifications', []) if not n.get('read', False))
    return jsonify({"count": unread})

# ---------- 评论 ----------
@app.route("/api/comments/<doc_id>", methods=["GET"])
def get_comments(doc_id):
    comments = load_comments(doc_id)
    users = load_users()
    for c in comments:
        uname = c.get("username")
        c["avatar"] = users[uname].get("avatar") if uname and uname in users else None
        c["role"] = users[uname].get("role", "user") if uname and uname in users else "user"
        c["is_dojin"] = users[uname].get("dojin", False) if uname and uname in users else False
    return jsonify(comments)

@app.route("/api/comments/<doc_id>", methods=["POST"])
@login_required
def add_comment(doc_id):
    username = g.username
    data = request.json
    selected_text = data.get("selectedText", "")
    comment_text = data.get("comment")
    anchor = data.get("anchor")
    parent_id = data.get("parentId")
    if not comment_text:
        return jsonify({"status": "error", "message": "评论内容不能为空"}), 400
    users = load_users()
    nickname = users.get(username, {}).get("nickname", username)
    avatar = users[username].get("avatar")
    comments = load_comments(doc_id)
    new_comment = {
        "id": str(uuid.uuid4()),
        "username": username,
        "nickname": nickname,
        "selectedText": selected_text,
        "comment": comment_text,
        "anchor": anchor,
        "parentId": parent_id,
        "timestamp": datetime.now().isoformat(),
        "likes": [],
        "avatar": avatar
    }
    comments.append(new_comment)
    # 顶级评论通知作者
    if not parent_id:
        meta = load_meta(doc_id)
        title = meta.get("title", doc_id)
        # 遍历所有作者，发送私信提醒
        for author_nick in meta.get("authors", []):
            # 根据昵称找到用户名
            author_username = None
            for uname, info in users.items():
                if info.get("nickname") == author_nick:
                    author_username = uname
                    break
            if author_username and author_username != username:
                # 发送私信给作者（WebSocket 已移除，消息已持久化）
                msgs = load_messages()
                msg_content = f"📝 有人评论了你的文章《{title}》：{comment_text[:50]}"
                new_msg = {
                    "id": str(uuid.uuid4()),
                    "from": username,
                    "to": author_username,
                    "content": msg_content,
                    "timestamp": datetime.now().isoformat(),
                    "read": False
                }
                msgs.append(new_msg)
                save_messages(msgs)
    if parent_id:
        for c in comments:
            if c["id"] == parent_id and c["username"] != username:
                # ===== 改为发送私信 =====
                msgs = load_messages()
                msg_content = f"💬 回复了你的评论：{comment_text[:50]}"
                new_msg = {
                    "id": str(uuid.uuid4()),
                    "from": username,
                    "to": c["username"],
                    "content": msg_content,
                    "timestamp": datetime.now().isoformat(),
                    "read": False
                }
                msgs.append(new_msg)
                save_messages(msgs)
                # WebSocket 已移除，消息已持久化到数据库
                break
    save_comments(doc_id, comments)
    return jsonify({"status": "ok", "comment": new_comment})

@app.route("/api/comments/<doc_id>/<comment_id>/like", methods=["POST"])
@login_required
def like_comment(doc_id, comment_id):
    username = g.username
    comments = load_comments(doc_id)
    users = load_users()
    nickname = users.get(username, {}).get("nickname", username)
    for c in comments:
        if c["id"] == comment_id:
            likes = c.get("likes", [])
            if username in likes:
                likes.remove(username)
                action = "unliked"
            else:
                likes.append(username)
                action = "liked"
            c["likes"] = likes
            save_comments(doc_id, comments)
            if action == "liked" and c["username"] != username:
                msgs = load_messages()
                msg_content = f"❤️ 赞了你的评论：{c['comment'][:50]}"
                new_msg = {
                    "id": str(uuid.uuid4()),
                    "from": username,
                    "to": c["username"],
                    "content": msg_content,
                    "timestamp": datetime.now().isoformat(),
                    "read": False
                }
                msgs.append(new_msg)
                save_messages(msgs)
                # WebSocket 已移除，消息已持久化到数据库

            # ✅ 必须在这里返回成功结果
            return jsonify({"status": "ok", "action": action, "likes": len(likes)})

    return jsonify({"status": "error", "message": "评论不存在"}), 404

@app.route("/api/comments/<doc_id>/<comment_id>", methods=["DELETE"])
@login_required
def delete_comment(doc_id, comment_id):
    username = g.username
    comments = load_comments(doc_id)
    new_comments = [c for c in comments if not (c["id"] == comment_id and c["username"] == username)]
    if len(new_comments) == len(comments):
        return jsonify({"status": "error", "message": "无权删除或评论不存在"}), 403
    save_comments(doc_id, new_comments)
    return jsonify({"status": "ok"})

# ---------- 头像 ----------
@app.route("/api/upload_avatar", methods=["POST"])
@login_required
def upload_avatar():
    username = g.username
    if 'avatar' not in request.files:
        return jsonify({"status": "error", "message": "没有文件"}), 400
    file = request.files['avatar']
    if file.filename == '':
        return jsonify({"status": "error", "message": "空文件名"}), 400

    # 大小检查
    file.seek(0, os.SEEK_END)
    size = file.tell()
    file.seek(0)
    if size > UPLOAD_MAX_SIZE:
        return jsonify({"status": "error", "message": f"文件过大，最大允许{UPLOAD_MAX_SIZE//1024//1024}MB"}), 400

    ext = file.filename.rsplit('.', 1)[-1].lower()
    if ext not in ['png', 'jpg', 'jpeg', 'gif', 'webp']:
        return jsonify({"status": "error", "message": "不支持的格式"}), 400
    safe_username = username.replace('/', '_').replace('\\', '_')
    filename = f"{safe_username}_{uuid.uuid4().hex[:8]}.{ext}"
    filepath = os.path.join(AVATARS_DIR, filename)
    file.save(filepath)
    optimize_image(filepath)

    # 头像缩略图（200x200）
    thumb_filename = f"{safe_username}_{uuid.uuid4().hex[:8]}_thumb.webp"
    thumb_filepath = os.path.join(AVATARS_DIR, thumb_filename)
    make_thumb(filepath, thumb_filepath, max_size=200)

    users = load_users()
    if username in users:
        old_avatar = users[username].get("avatar")
        old_thumb = users[username].get("avatar_thumb")
        if old_avatar:
            old_path = os.path.join(AVATARS_DIR, old_avatar)
            if os.path.exists(old_path):
                os.remove(old_path)
        if old_thumb:
            old_thumb_path = os.path.join(AVATARS_DIR, old_thumb)
            if os.path.exists(old_thumb_path):
                os.remove(old_thumb_path)
        users[username]["avatar"] = filename
        users[username]["avatar_thumb"] = thumb_filename
        save_users(users)
    return jsonify({"status": "ok", "avatar": filename})

@app.route("/avatars/<filename>")
def serve_avatar(filename):
    return send_from_directory(AVATARS_DIR, filename)

# ---------- 信件（站内私信）API ----------
@app.route("/api/messages/inbox", methods=["GET"])
@login_required
def inbox():
    """合并通知与私信列表"""
    current = g.username
    users = load_users()

    # 1. 获取通知（复用已有逻辑，但只取前20条，避免列表过长）
    user_data = users.get(current, {})
    notifications = user_data.get('notifications', [])[:20]
    noti_list = []
    for n in notifications:
        noti_list.append({
            "id": n["id"],
            "type": "notification",
            "kind": n.get("type"),            # reply / like
            "from_user": n.get("from_user"),
            "from_avatar": n.get("from_avatar"),
            "content": n.get("content", ""),
            "doc_id": n.get("doc_id"),
            "comment_id": n.get("comment_id"),
            "timestamp": n.get("timestamp"),
            "read": n.get("read", False)
        })

    # 2. 获取私信对话列表（复用已有的 conversations 逻辑）
    all_msgs = load_messages()
    conv_dict = {}
    for m in all_msgs:
        if m["from"] != current and m["to"] != current:
            continue
        other = m["from"] if m["to"] == current else m["to"]
        if other not in conv_dict or m["timestamp"] > conv_dict[other]["timestamp"]:
            conv_dict[other] = m
    msg_list = []
    for other, last_msg in conv_dict.items():
        other_info = users.get(other)
        unread_count = sum(1 for m in all_msgs if m["to"] == current and m["from"] == other and not m.get("read", False))
        msg_list.append({
            "id": last_msg["id"],
            "type": "message",
            "with_user": other,
            "nickname": other_info["nickname"] if other_info else other,
            "avatar": other_info.get("avatar") if other_info else None,
            "last_content": last_msg["content"][:100],
            "last_timestamp": last_msg["timestamp"],
            "unread_count": unread_count,
            "read": last_msg.get("read", False)
        })

    # 3. 合并并按时间倒序
    combined = noti_list + msg_list
    combined.sort(key=lambda x: x.get("last_timestamp") or x.get("timestamp"), reverse=True)

    return jsonify(combined)

@app.route("/api/read_all", methods=["POST"])
@login_required
def read_all():
    """标记所有通知和私信为已读"""
    username = g.username
    # 标记通知
    users = load_users()
    if username in users:
        for n in users[username].get('notifications', []):
            n['read'] = True
        save_users(users)
    # 标记私信
    msgs = load_messages()
    updated = False
    for m in msgs:
        if m['to'] == username and not m.get('read', False):
            m['read'] = True
            updated = True
    if updated:
        save_messages(msgs)
    return jsonify({"status": "ok"})

@app.route("/api/messages", methods=["GET"])
@login_required
def get_messages():
    """获取当前用户与指定用户的完整对话历史"""
    current = g.username
    with_user = request.args.get("with", "").strip()
    if not with_user:
        return jsonify({"status": "error", "message": "缺少参数"}), 400
    users = load_users()
    if with_user not in users:
        return jsonify({"status": "error", "message": "用户不存在"}), 404
    msgs = load_messages()
    # 筛选双方的消息
    conv = [m for m in msgs if (m["from"] == current and m["to"] == with_user) or (m["from"] == with_user and m["to"] == current)]
    conv.sort(key=lambda x: x["timestamp"])
    result = []
    for m in conv:
        from_user = users.get(m["from"])
        to_user = users.get(m["to"])
        result.append({
            "id": m["id"],
            "from": m["from"],
            "from_nickname": from_user["nickname"] if from_user else m["from"],
            "from_avatar": from_user.get("avatar") if from_user else None,
            "to": m["to"],
            "to_nickname": to_user["nickname"] if to_user else m["to"],
            "to_avatar": to_user.get("avatar") if to_user else None,
            "content": m["content"],
            "timestamp": m["timestamp"],
            "read": m.get("read", False)
        })
    return jsonify(result)

@app.route("/api/messages/conversations", methods=["GET"])
@login_required
def get_conversations():
    """获取当前用户的所有聊天列表（含最后一条消息、未读数）"""
    username = g.username
    msgs = load_messages()
    users = load_users()
    conversations = {}
    for m in msgs:
        if m["from"] != username and m["to"] != username:
            continue
        other = m["from"] if m["to"] == username else m["to"]
        if other not in conversations:
            conversations[other] = []
        conversations[other].append(m)
    result = []
    for other, msgs_list in conversations.items():
        msgs_list.sort(key=lambda x: x["timestamp"])
        last = msgs_list[-1]
        unread_count = sum(1 for m in msgs_list if m["to"] == username and not m.get("read", False))
        other_info = users.get(other)
        result.append({
            "with_user": other,
            "nickname": other_info["nickname"] if other_info else other,
            "avatar": other_info.get("avatar") if other_info else None,
            "last_message": {
                "content": last["content"][:100],
                "timestamp": last["timestamp"],
                "from_me": last["from"] == username,
                "read": last.get("read", False)
            },
            "unread_count": unread_count
        })
    result.sort(key=lambda x: x["last_message"]["timestamp"], reverse=True)
    return jsonify(result)

@app.route("/api/messages", methods=["POST"])
@login_required
def send_message():
    """发送私信，并同时推送给在线用户 + 写入通知"""
    current = g.username
    data = request.json
    to_user = data.get("to", "").strip()
    content = data.get("content", "").strip()
    if not to_user or not content:
        return jsonify({"status": "error", "message": "缺少参数"}), 400
    users = load_users()
    if to_user not in users:
        return jsonify({"status": "error", "message": "用户不存在"}), 404
    if to_user == current:
        return jsonify({"status": "error", "message": "不能给自己发消息"}), 400
    msgs = load_messages()
    new_msg = {
        "id": str(uuid.uuid4()),
        "from": current,
        "to": to_user,
        "content": content,
        "timestamp": datetime.now().isoformat(),
        "read": False
    }
    msgs.append(new_msg)
    save_messages(msgs)
    # WebSocket 已移除，消息已持久化到数据库
    return jsonify({"status": "ok", "message": new_msg})

@app.route("/api/messages/read", methods=["POST"])
@login_required
def mark_messages_read():
    """将当前用户与指定用户的对话标记为已读"""
    current = g.username
    data = request.json
    with_user = data.get("with", "").strip()
    if not with_user:
        return jsonify({"status": "error", "message": "缺少参数"}), 400
    msgs = load_messages()
    updated = False
    for m in msgs:
        if m["to"] == current and m["from"] == with_user and not m.get("read", False):
            m["read"] = True
            updated = True
    if updated:
        save_messages(msgs)
    return jsonify({"status": "ok"})

@app.route("/api/messages/unread_count", methods=["GET"])
@login_required
def unread_message_count():
    """获取当前用户所有未读私信总数"""
    current = g.username
    msgs = load_messages()
    count = sum(1 for m in msgs if m["to"] == current and not m.get("read", False))
    return jsonify({"count": count})

@app.route("/api/messages/<msg_id>", methods=["DELETE"])
@login_required
def delete_message(msg_id):
    """删除自己发送的一条私信"""
    current = g.username
    msgs = load_messages()
    new_msgs = []
    deleted = False
    for m in msgs:
        if m["id"] == msg_id and m["from"] == current:
            deleted = True
            continue
        new_msgs.append(m)
    if not deleted:
        return jsonify({"status": "error", "message": "无权删除或消息不存在"}), 403
    save_messages(new_msgs)
    return jsonify({"status": "ok"})

# ===== 表情包 / stickers =====
@app.route("/api/stickers", methods=["GET"])
@login_required
def get_stickers():
    """获取当前用户上传的表情包"""
    username = g.username
    user_sticker_dir = os.path.join(STICKERS_DIR, username)
    os.makedirs(user_sticker_dir, exist_ok=True)
    stickers = []
    for fname in sorted(os.listdir(user_sticker_dir)):
        if fname.lower().endswith((".png", ".jpg", ".jpeg", ".gif", ".webp")):
            stickers.append({
                "url": f"/static/uploads/stickers/{username}/{fname}",
                "name": fname
            })
    return jsonify({"stickers": stickers})

@app.route("/api/stickers/upload", methods=["POST"])
@login_required
def upload_sticker():
    """上传表情包"""
    username = g.username
    if "sticker" not in request.files:
        return jsonify({"status": "error", "message": "未选择文件"}), 400
    file = request.files["sticker"]
    if not file.filename:
        return jsonify({"status": "error", "message": "空文件"}), 400
    size = 0
    try:
        size = os.fstat(file.fileno()).st_size
    except Exception:
        file.seek(0, os.SEEK_END)
        size = file.tell()
        file.seek(0)
    if size > 5 * 1024 * 1024:
        return jsonify({"status": "error", "message": "文件过大，最大5MB"}), 400
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else "png"
    if ext not in ("png", "jpg", "jpeg", "gif", "webp"):
        return jsonify({"status": "error", "message": "不支持的格式"}), 400
    filename = f"{uuid.uuid4().hex[:12]}.{ext}"
    user_sticker_dir = os.path.join(STICKERS_DIR, username)
    os.makedirs(user_sticker_dir, exist_ok=True)
    filepath = os.path.join(user_sticker_dir, filename)
    file.save(filepath)
    return jsonify({
        "status": "ok",
        "url": f"/static/uploads/stickers/{username}/{filename}",
        "name": filename
    })

@app.route("/api/stickers/delete", methods=["POST"])
@login_required
def delete_sticker():
    """删除表情包"""
    username = g.username
    data = request.get_json()
    name = data.get("name", "")
    if not name:
        return jsonify({"status": "error", "message": "参数缺失"}), 400
    safe_name = os.path.basename(name)
    filepath = os.path.join(STICKERS_DIR, username, safe_name)
    if os.path.exists(filepath):
        os.remove(filepath)
        return jsonify({"status": "ok"})
    return jsonify({"status": "error", "message": "文件不存在"}), 404

# ──── 大纲系统已删除 ────
if __name__ == "__main__":
    from waitress import serve
    serve(app, host="0.0.0.0", port=5000)